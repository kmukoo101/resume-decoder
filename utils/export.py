from docx import Document
from io import BytesIO

def export_to_docx(title: str, sections: dict):
    """Generate a DOCX file from a dictionary of resume sections."""
    doc = Document()
    doc.add_heading(title, 0)

    for section, content in sections.items():
        doc.add_heading(section, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
